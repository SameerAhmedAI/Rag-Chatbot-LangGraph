"""
DOCX document loader (concrete Strategy).
Extracts paragraph text and table content, preserving basic structure.
"""

from langchain_core.documents import Document
import docx

from app.ingestion.base_loader import DocumentLoader


class DocxLoader(DocumentLoader):
    """Loads .docx files into a single Document containing paragraphs + tables."""

    @property
    def file_type(self) -> str:
        return "docx"

    def load(self, file_path: str, source_name: str) -> list[Document]:
        doc = docx.Document(file_path)
        parts: list[str] = []

        # Paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                parts.append(para.text.strip())

        # Tables (flatten rows into readable lines)
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells)
                if row_text.strip(" |"):
                    parts.append(row_text)

        full_text = "\n".join(parts)

        if not full_text.strip():
            return []

        return [
            Document(
                page_content=full_text,
                metadata={
                    "source": source_name,
                    "file_type": self.file_type,
                },
            )
        ]